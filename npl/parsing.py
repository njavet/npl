from pathlib import Path
import re
from docx import Document
from docx.oxml.ns import qn
from langchain.text_splitter import RecursiveCharacterTextSplitter

# project imports
from mars.conf import conf


def get_doc_sections(docx_path: Path) -> list[str]:
    text = extract_text_from_docx(docx_path)
    parsed_text = parse_text_to_llm_input(text)
    sections = split_text(parsed_text)
    sections = split_big_sections(sections)
    sections = unify_small_sections(sections)
    return sections


def unify_small_sections(sections: list[str]) -> list[str]:
    new_sections = []
    curr = ''
    for section in sections:
        if len(curr) + len(section) <= conf.DOC_CHUNK_SIZE:
            curr = '\n\n'.join([curr, section])
        else:
            new_sections.append(curr)
            curr = ''
    return new_sections


def split_big_sections(sections: list[str]) -> list[str]:
    # TODO how to parse a docx file into desired llm input style ?
    #  TODO extract section header
    for section in sections:
        if len(section) > conf.DOC_CHUNK_SIZE:
            # TODO split
            pass
    return sections


# TODO if a section separated with '\n\n' contains too many characters,
#  split into two sections with same title
def parse_text_to_llm_input(text: str) -> str:
    # remove leading and training ws
    text = '\n'.join(line.strip() for line in text.splitlines())
    # replace tabs
    text = re.sub(r'\t+', ' ', text)
    # replace multiple ws with one ws
    text = re.sub(r'[ ]{2,}', ' ', text)
    # replace 3+ newlines with 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # replace non semantic newlines
    text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
    # restore bullet point list
    text = re.sub(r'(?<!\n)\* ', r'\n* ', text)
    text = '\n'.join(line.strip() for line in text.splitlines())
    return text


def extract_docx_text(docx_dir: Path = conf.DOCX_DIR) -> None:
    for docx_path in docx_dir.glob('*.docx'):
        text = extract_text_from_docx(docx_path)
        out_file = (conf.TEXT_DIR / docx_path.stem).with_suffix('.txt')
        with open(out_file, 'w') as f:
            f.write(text)


def extract_text_from_docx(docx_path: Path) -> str:
    doc = Document(docx_path)
    return '\n'.join([p.text.strip() for p in doc.paragraphs])


def split_text(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=conf.DOC_CHUNK_SIZE,
        chunk_overlap=0,
        separators=conf.DOC_SEPARATORS
    )
    chunks = splitter.split_text(text)
    return chunks


def split_docx(dox_path: Path) -> list[str]:
    text = extract_text_from_docx(dox_path)
    return split_text(text)


def clean_medical_body(doc) -> dict[str, list[str]]:
    def strip_headers_footers() -> None:
        for sect in doc.sections:
            for tag in ('headerReference', 'footerReference'):
                for ref in sect._sectPr.findall(qn(f'w:{tag}')):
                    sect._sectPr.remove(ref)
            for part in (
                    sect.header, sect.first_page_header, sect.even_page_header,
                    sect.footer, sect.first_page_footer, sect.even_page_footer
            ):
                part._element.clear()

    strip_headers_footers()

    out: dict[str, list[str]] = {}
    current = None

    for p in doc.paragraphs:
        if p.part is not doc.part:
            continue
        text = p.text.strip()
        if not text:
            continue

        # new section ?
        if text in conf.ALLOWED_HEADINGS:
            current = text
            out[current] = []
            continue

        if current:
            out[current].append(text)

    return {k: v for k, v in out.items() if v}
